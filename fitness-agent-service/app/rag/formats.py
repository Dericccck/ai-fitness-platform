"""将多格式文档解析为保留结构的 RAG 内容块。"""

from __future__ import annotations

import re
import unicodedata
from collections import Counter
from collections.abc import Sequence
from dataclasses import dataclass, replace
from io import BytesIO
from itertools import pairwise
from pathlib import PurePosixPath
from typing import Literal, Protocol

from .text import clean_markdown

# 标准解析块类型：TEXT 普通文本；TABLE 保留表头和行范围的结构化表格文本。
BlockKind = Literal["TEXT", "TABLE"]

# PDF 页面路由：NORMAL 正常解析；OCR_REQUIRED 缺少文字层、必须 OCR；
# VISUAL_REVIEW_REQUIRED 图片承载动作/姿态信息、必须人工视觉审核；
# OCR_AND_VISUAL_REVIEW_REQUIRED 两项都不能绕过，人工审核也不能替代 OCR。
PdfPageRoute = Literal[
    "NORMAL",
    "OCR_REQUIRED",
    "VISUAL_REVIEW_REQUIRED",
    "OCR_AND_VISUAL_REVIEW_REQUIRED",
]


class UnsupportedDocumentFormatError(ValueError):
    """入库解析器注册表没有适用于该来源格式的安全解析器。"""


class DocumentParseError(ValueError):
    """支持的文件无法解析，或不包含可用内容。"""


@dataclass(frozen=True)
class ParsedBlock:
    """保留来源坐标的标准文本/表格内容块。"""

    kind: BlockKind
    content: str
    heading_path: tuple[str, ...] = ()
    source_page: int | None = None
    source_sheet: str | None = None
    table_index: int | None = None
    row_start: int | None = None
    row_end: int | None = None
    metadata: dict[str, str | int | bool] | None = None
    # PDF 页面可能把一个章节拆成多个文本块。保留章节级父上下文后，后续子块可以
    # 共享完整标题和邻近正文，避免检索命中一句孤立文本时失去“这是哪个动作/原则”的语义。
    parent_content: str | None = None


@dataclass(frozen=True)
class PdfPageProfile:
    """PDF 单页在进入 OCR、人工审核和 Embedding 前的可审计画像。

    画像只使用 PDF 自身的对象坐标和原生文字层，不调用视觉模型，也不尝试判断
    健身动作是否标准。图片密集页只会被安全地路由到专业审核，避免模型在没有
    教练或医疗人员确认时生成动作、伤病或禁忌结论。
    """

    page_number: int
    image_count: int
    image_area_ratio: float
    native_text_chars: int
    text_area_ratio: float
    table_count: int
    caption_count: int
    route: PdfPageRoute
    reasons: tuple[str, ...] = ()
    # 下面这些字段是版面解析的证据，不参与权限判断。它们让人工审核和离线质量
    # 报告能够回答“为什么这一页被这样处理”，也避免把排版猜测伪装成正文内容。
    detected_columns: int = 1
    removed_repeated_edge_lines: int = 0
    toc_detected: bool = False
    dehyphenated_line_breaks: int = 0

    def as_dict(self) -> dict[str, object]:
        """输出稳定审计结构，供离线报告和后续审核 API 复用。"""

        return {
            "page_number": self.page_number,
            "image_count": self.image_count,
            "image_area_ratio": round(self.image_area_ratio, 6),
            "native_text_chars": self.native_text_chars,
            "text_area_ratio": round(self.text_area_ratio, 6),
            "table_count": self.table_count,
            "caption_count": self.caption_count,
            "route": self.route,
            "reasons": list(self.reasons),
            "detected_columns": self.detected_columns,
            "removed_repeated_edge_lines": self.removed_repeated_edge_lines,
            "toc_detected": self.toc_detected,
            "dehyphenated_line_breaks": self.dehyphenated_line_breaks,
        }


@dataclass(frozen=True)
class ParsedDocument:
    """分块和 Embedding 前的解析器输出及页面级处理证据。"""

    blocks: tuple[ParsedBlock, ...]
    media_type: str
    warnings: tuple[str, ...] = ()
    page_profiles: tuple[PdfPageProfile, ...] = ()


@dataclass(frozen=True)
class PdfPageRoutingPolicy:
    """企业部署可调整的 PDF 页面路由阈值。

    默认值偏保守：只要大面积图片可能承载动作示范，而原生文字层不足以完整解释
    页面，就要求人工视觉审核。阈值必须由真实健身资料评测后再调整，不能由上传者
    或 LLM 在请求中覆盖。
    """

    min_image_area_ratio: float = 0.45
    max_image_page_text_chars: int = 600
    max_image_page_text_area_ratio: float = 0.25
    min_ocr_text_chars: int = 12

    def __post_init__(self) -> None:
        if not 0 <= self.min_image_area_ratio <= 1:
            raise ValueError("min_image_area_ratio must be between 0 and 1")
        if self.max_image_page_text_chars < 0:
            raise ValueError("max_image_page_text_chars must not be negative")
        if not 0 <= self.max_image_page_text_area_ratio <= 1:
            raise ValueError("max_image_page_text_area_ratio must be between 0 and 1")
        if self.min_ocr_text_chars < 0:
            raise ValueError("min_ocr_text_chars must not be negative")


class DocumentParser(Protocol):
    """每种支持的文件格式都需要实现的解析器契约。"""

    def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        """解析字节内容，不将不可信上传文件写入本地路径。"""


@dataclass(frozen=True)
class _PdfTextLine:
    """PDF 原生文字层的一行及其版面坐标。

    ``pdfplumber.extract_text`` 只返回字符串，会丢掉列位置；这也是演示文稿和双栏
    指南出现“左列一句、右列一句交叉拼接”的根因。内部先保留坐标，确认页面确实是
    多区域版式后才启用坐标排序，普通单栏页面继续使用 pdfplumber 的成熟文本流。
    """

    text: str
    x0: float
    x1: float
    top: float
    bottom: float

    @property
    def width(self) -> float:
        """返回文字行的横向跨度。"""

        return max(0.0, self.x1 - self.x0)


class PdfOcrProvider(Protocol):
    """扫描型 PDF 的 OCR 边界，具体实现属于独立 OCR 服务。"""

    def parse(
        self,
        content: bytes,
        *,
        file_name: str,
        pages: Sequence[int] = (),
    ) -> ParsedDocument:
        """返回指定 PDF 页面中保留结构的 OCR 内容块。"""


class DocumentParserRegistry:
    """根据标准化扩展名选择解析器，并执行上传大小限制。"""

    def __init__(
        self,
        *,
        max_source_bytes: int = 20 * 1024 * 1024,
        pdf_ocr_provider: PdfOcrProvider | None = None,
        pdf_page_routing_policy: PdfPageRoutingPolicy | None = None,
    ) -> None:
        self.max_source_bytes = max_source_bytes
        self._parsers: dict[str, DocumentParser] = {
            ".md": MarkdownParser(),
            ".markdown": MarkdownParser(),
            ".txt": MarkdownParser(),
            ".pdf": PdfParser(
                ocr_provider=pdf_ocr_provider,
                routing_policy=pdf_page_routing_policy,
            ),
            ".docx": DocxParser(),
            ".xlsx": XlsxParser(),
        }

    def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        """解析支持的扩展名，并尽早拒绝超大或空上传文件。"""

        if not content:
            raise DocumentParseError("document content must not be empty")
        if len(content) > self.max_source_bytes:
            raise DocumentParseError("document exceeds the configured size limit")
        suffix = PurePosixPath(file_name).suffix.lower()
        parser = self._parsers.get(suffix)
        if parser is None:
            supported = ", ".join(sorted(self._parsers))
            raise UnsupportedDocumentFormatError(
                f"unsupported document extension {suffix or '<none>'}; supported: {supported}"
            )
        try:
            parsed = parser.parse(content, file_name=file_name)
        except (DocumentParseError, UnsupportedDocumentFormatError):
            raise
        except Exception as exc:
            raise DocumentParseError(f"failed to parse {file_name}") from exc
        # 待 OCR/视觉审核的 PDF 即使暂时没有可索引正文，也必须保留页面画像供
        # 质量报告和人工审核使用。真正发布时由入库门禁 fail-closed；其他格式
        # 仍然要求至少产生一个内容块，避免空文档进入任务队列。
        if not parsed.blocks and not any(
            profile.route != "NORMAL" for profile in parsed.page_profiles
        ):
            raise DocumentParseError(f"document {file_name} contains no indexable content")
        return parsed


class MarkdownParser:
    """解析 Markdown/文本，同时保留现有的标题感知分块路径。"""

    def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        try:
            raw = content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("Markdown/text must be UTF-8") from exc
        cleaned = clean_markdown(raw)
        return ParsedDocument(
            blocks=(ParsedBlock(kind="TEXT", content=cleaned),),
            media_type="text/markdown"
            if file_name.lower().endswith((".md", ".markdown"))
            else "text/plain",
        )


class PdfParser:
    """提取 PDF 页面文本和表格，并在进入分块前完成一轮保守清洗。

    PDF 的文本层经常把网页导航、页眉页脚和字体映射字符混入正文。这里先按页收集
    原始行，再做模板过滤、重复行识别、Unicode 规范化和有限的段落重组，最后才生成
    ``ParsedBlock``。这样父节点不会直接继承整页的版式噪声；原始文件仍由对象存储或
    本地资料目录保留，清洗规则变化后可以重新解析。
    """

    def __init__(
        self,
        *,
        ocr_provider: PdfOcrProvider | None = None,
        routing_policy: PdfPageRoutingPolicy | None = None,
    ) -> None:
        self.ocr_provider = ocr_provider
        self.routing_policy = routing_policy or PdfPageRoutingPolicy()

    def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        import pdfplumber

        blocks: list[ParsedBlock] = []
        warnings: list[str] = []
        missing_pages: list[int] = []
        page_profiles: list[PdfPageProfile] = []
        with pdfplumber.open(BytesIO(content)) as document:
            page_records: list[
                tuple[int, list[str], list[list[list[str | None]]], PdfPageProfile]
            ] = []
            for page_number, page in enumerate(document.pages, start=1):
                tables = _extract_pdf_tables(page)
                raw_lines, detected_columns = _extract_pdf_text_lines(page, tables)
                raw_text = "\n".join(raw_lines)
                profile = _profile_pdf_page(
                    page,
                    page_number=page_number,
                    raw_text=raw_text,
                    tables=tables,
                    policy=self.routing_policy,
                )
                profile = replace(profile, detected_columns=detected_columns)
                page_records.append((page_number, raw_lines, tables, profile))
                page_profiles.append(profile)

            repeated_lines = _repeated_pdf_lines([raw_lines for _, raw_lines, _, _ in page_records])
            for page_number, raw_lines, tables, profile in page_records:
                cleaned_lines = _clean_pdf_lines(raw_lines, repeated_lines=repeated_lines)
                toc_detected = _is_pdf_toc_page(cleaned_lines)
                if toc_detected:
                    cleaned_lines = []
                removed_repeated_edge_lines = sum(
                    _normalize_pdf_line(line) in repeated_lines for line in raw_lines
                )
                profile = replace(
                    profile,
                    removed_repeated_edge_lines=removed_repeated_edge_lines,
                    toc_detected=toc_detected,
                    # 断词修复发生在“行合并”阶段，先统计证据，再把统计值写入每个
                    # 页面画像。这里不把普通连字符替换成空格，避免损坏 compound
                    # words、编号和 URL。
                    dehyphenated_line_breaks=_count_pdf_dehyphenations(cleaned_lines),
                )
                page_profiles[page_number - 1] = profile
                page_metadata: dict[str, str | int | bool] = {
                    "parser": "pdfplumber",
                    "cleaned": True,
                    "page_route": profile.route,
                    "image_count": profile.image_count,
                    "image_area_basis_points": round(profile.image_area_ratio * 10_000),
                    "native_text_chars": profile.native_text_chars,
                    "text_area_basis_points": round(profile.text_area_ratio * 10_000),
                    "caption_count": profile.caption_count,
                    "detected_columns": profile.detected_columns,
                    "removed_repeated_edge_lines": profile.removed_repeated_edge_lines,
                    "toc_detected": profile.toc_detected,
                    "dehyphenated_line_breaks": profile.dehyphenated_line_breaks,
                }
                for text_block, heading_path, parent_content in _pdf_text_blocks_with_context(
                    cleaned_lines
                ):
                    blocks.append(
                        ParsedBlock(
                            kind="TEXT",
                            content=text_block,
                            heading_path=heading_path,
                            source_page=page_number,
                            metadata=page_metadata,
                            parent_content=parent_content,
                        )
                    )
                for table_index, table in enumerate(tables):
                    normalized = _table_to_markdown(table)
                    if normalized:
                        blocks.append(
                            ParsedBlock(
                                kind="TABLE",
                                content=normalized,
                                source_page=page_number,
                                table_index=table_index,
                                row_start=1,
                                row_end=len(table),
                                metadata={
                                    **page_metadata,
                                    "table_header_repeated": True,
                                    "table_column_count": _table_column_count(table),
                                    "table_row_count": _table_row_count(table),
                                    "table_header_signature": _table_header_signature(table),
                                    "table_header_key": _table_header_key(table),
                                    "table_continuation_status": "SINGLE_PAGE",
                                },
                            )
                        )
                if profile.route in {"OCR_REQUIRED", "OCR_AND_VISUAL_REVIEW_REQUIRED"}:
                    missing_pages.append(page_number)
        if missing_pages and self.ocr_provider is not None:
            ocr_document = self.ocr_provider.parse(
                content,
                file_name=file_name,
                pages=tuple(missing_pages),
            )
            blocks.extend(ocr_document.blocks)
            warnings.extend(ocr_document.warnings)
            resolved_pages = {
                block.source_page
                for block in ocr_document.blocks
                if block.source_page is not None and block.content.strip()
            }
            # OCR 必须逐页返回来源页码；只有确实返回了可用内容的页面才解除阻断。
            # 未返回、返回空块或缺少页码的页面继续保持 OCR_REQUIRED，防止静默漏页。
            resolved_profiles: list[PdfPageProfile] = []
            for profile in page_profiles:
                if profile.page_number not in resolved_pages:
                    resolved_profiles.append(profile)
                elif profile.route == "OCR_REQUIRED":
                    resolved_profiles.append(
                        replace(
                            profile,
                            route="NORMAL",
                            reasons=profile.reasons + ("OCR_COMPLETED",),
                        )
                    )
                elif profile.route == "OCR_AND_VISUAL_REVIEW_REQUIRED":
                    # OCR 只补齐图片中的文字，不能证明健身动作、姿态和风险已经由
                    # 教练或医疗专业人员审核，因此组合状态只解除 OCR 部分。
                    resolved_profiles.append(
                        replace(
                            profile,
                            route="VISUAL_REVIEW_REQUIRED",
                            reasons=profile.reasons + ("OCR_COMPLETED",),
                        )
                    )
                else:
                    resolved_profiles.append(profile)
            page_profiles = resolved_profiles
        # 只根据 OCR 调用后的最终路由生成警告，避免 OCR 已成功的页面继续携带
        # “requires OCR”旧警告。OCR 服务自身的低置信度等警告仍然完整保留。
        for profile in page_profiles:
            if profile.route == "OCR_REQUIRED":
                warnings.append(f"page {profile.page_number} requires OCR before publication")
            elif profile.route == "VISUAL_REVIEW_REQUIRED":
                warnings.append(
                    f"page {profile.page_number} is image-heavy and requires professional visual review"
                )
            elif profile.route == "OCR_AND_VISUAL_REVIEW_REQUIRED":
                warnings.append(
                    f"page {profile.page_number} requires OCR and professional visual review"
                )
        blocks = _annotate_pdf_table_continuations(blocks)
        return ParsedDocument(
            tuple(blocks),
            "application/pdf",
            tuple(warnings),
            tuple(page_profiles),
        )


_PDF_TOC_MARKERS = ("table of contents", "目录")
_PDF_DOT_LEADER = re.compile(r"(?:\.{4,}|·{4,}|…{2,})\s*\d+\s*$")
_PDF_TEMPLATE_PATTERNS = (
    re.compile(r"无障碍浏览\s*网站导航\s*公务员邮箱"),
    re.compile(r"请输入关键字\s*搜"),
    re.compile(r"^(?:首\s*页|首页)\s*[>＞].*$"),
    re.compile(r"^首\s*页\s+机\s*构\s+公\s*开\s+资\s*讯\s+服\s*务\s+互\s*动$"),
    re.compile(r"(?:打印此页|关闭窗口|关闭窗⼜|print this page|close window)", re.IGNORECASE),
    re.compile(r"(?:版权所有|网站标识码|ICP备|公安网安备|通讯地址|邮政编码|信访电话)"),
    re.compile(r"^联系电话[:：]|^字体[:：]\s*[大中小]+$|^发布时间[:：]"),
    re.compile(r"^(?:accessibility|site navigation|copyright)\b", re.IGNORECASE),
    # HHS 演示文稿每页都会重复这段来源声明；它是页脚，不是可检索知识。
    re.compile(
        r"^information adapted from the physical activity guidelines for americans,?\s*"
        r"2nd edition\.?(?:\s+available at .*)?$",
        re.IGNORECASE,
    ),
    # 正式指南的页脚会随着页码变化，不能只依赖“跨页完全相同”识别。
    re.compile(
        r"^(?:\d+\s+)?physical activity guidelines for americans(?:\s*\|.*)?$",
        re.IGNORECASE,
    ),
    re.compile(r"^physical activity guidelines for americans\s*\|\s*summary\s+\d+$", re.IGNORECASE),
)
_PDF_SENTENCE_END = frozenset("。！？.!?；;：:")
_CJK = re.compile(r"[\u3400-\u9fff]")
_PDF_ENGLISH_HEADING_WORDS = frozenset(
    ["and", "or", "the", "of", "for", "to", "in", "on", "with", "from", "among"]
)
_PDF_COMMON_COMPOUNDS = frozenset(
    {
        "activity-related",
        "bone-strengthening",
        "care-provider",
        "cardiorespiratory-fitness",
        "disease-specific",
        "evidence-based",
        "fall-related",
        "health-care",
        "health-related",
        "long-term",
        "moderate-to-vigorous",
        "muscle-strengthening",
        "self-reported",
    }
)
_PDF_CAPTION = re.compile(
    r"^(?:图\s*\d+(?:[-–.]\d+)*|figure\s*\d+(?:[-–.]\d+)*|fig\.\s*\d+)",
    re.IGNORECASE,
)
_PDF_FULL_WIDTH_RATIO = 0.60


def _profile_pdf_page(
    page: object,
    *,
    page_number: int,
    raw_text: str,
    tables: Sequence[Sequence[Sequence[str | None]]],
    policy: PdfPageRoutingPolicy,
) -> PdfPageProfile:
    """根据页面原生对象生成确定性画像，不依赖 Linux OCR 或外部模型。

    图片面积和文字面积使用矩形并集，而不是简单累加。PDF 常把一张图拆成多个重叠
    对象；直接累加会让比例超过 100%，并把普通页误判成图片密集页。坐标异常的对象
    被忽略，最终比例始终限制在 ``0..1``。
    """

    width = _positive_number(getattr(page, "width", 0.0))
    height = _positive_number(getattr(page, "height", 0.0))
    page_area = width * height
    image_rectangles = _pdf_object_rectangles(getattr(page, "images", ()), width, height)
    text_rectangles = _pdf_object_rectangles(getattr(page, "chars", ()), width, height)
    image_area_ratio = _area_ratio(_rectangle_union_area(image_rectangles), page_area)
    text_area_ratio = _area_ratio(_rectangle_union_area(text_rectangles), page_area)
    native_text_chars = len(re.sub(r"\s+", "", raw_text))
    caption_count = sum(
        bool(_PDF_CAPTION.match(_normalize_pdf_line(line))) for line in raw_text.splitlines()
    )

    reasons: list[str] = []
    ocr_required = native_text_chars < policy.min_ocr_text_chars and not tables
    # “图片多、文字少”要求字符数和文字覆盖面积同时偏低。仅使用 OR 会把带整页
    # 扫描底图、但已有完整 OCR 文字层的普通医学指南误判成动作图片。
    # 完全没有原生文字的页面只能先证明需要 OCR；OCR/版面识别完成后才能判断图片
    # 是否还承载姿态或动作信息。存在少量原生说明时则保守保留组合审核状态。
    visual_review_required = (
        image_area_ratio >= policy.min_image_area_ratio
        and native_text_chars <= policy.max_image_page_text_chars
        and text_area_ratio <= policy.max_image_page_text_area_ratio
        and (not ocr_required or native_text_chars > 0)
    )
    if ocr_required:
        reasons.append("NATIVE_TEXT_INSUFFICIENT")
        if image_rectangles:
            reasons.append("PAGE_CONTAINS_IMAGES")
    if visual_review_required:
        reasons.append("IMAGE_AREA_HIGH")
        if native_text_chars <= policy.max_image_page_text_chars:
            reasons.append("TEXT_DENSITY_LOW")
        if caption_count:
            reasons.append("IMAGE_CAPTION_PRESENT")

    if ocr_required and visual_review_required:
        route: PdfPageRoute = "OCR_AND_VISUAL_REVIEW_REQUIRED"
    elif ocr_required:
        route = "OCR_REQUIRED"
    elif visual_review_required:
        route = "VISUAL_REVIEW_REQUIRED"
    else:
        route = "NORMAL"

    return PdfPageProfile(
        page_number=page_number,
        image_count=len(image_rectangles),
        image_area_ratio=image_area_ratio,
        native_text_chars=native_text_chars,
        text_area_ratio=text_area_ratio,
        table_count=len(tables),
        caption_count=caption_count,
        route=route,
        reasons=tuple(reasons),
    )


def _pdf_object_rectangles(
    objects: object,
    page_width: float,
    page_height: float,
) -> list[tuple[float, float, float, float]]:
    """读取 pdfplumber 图片/字符坐标并裁剪到页面范围。"""

    if not isinstance(objects, Sequence):
        return []
    rectangles: list[tuple[float, float, float, float]] = []
    for item in objects:
        if not isinstance(item, dict):
            continue
        x0 = _number(item.get("x0"))
        x1 = _number(item.get("x1"))
        top = _number(item.get("top"))
        bottom = _number(item.get("bottom"))
        if None in (x0, x1, top, bottom):
            continue
        assert x0 is not None and x1 is not None and top is not None and bottom is not None
        left = min(max(min(x0, x1), 0.0), page_width)
        right = min(max(max(x0, x1), 0.0), page_width)
        upper = min(max(min(top, bottom), 0.0), page_height)
        lower = min(max(max(top, bottom), 0.0), page_height)
        if right > left and lower > upper:
            rectangles.append((left, upper, right, lower))
    return rectangles


def _rectangle_union_area(rectangles: Sequence[tuple[float, float, float, float]]) -> float:
    """通过扫描线计算矩形并集面积，避免重叠对象重复计数。

    事件只维护当前横向切片内的纵向区间。相比“每个 x 区间重新扫描全部字符”，
    该实现不会在长 PDF 的数千文字对象上产生明显的平方级开销。
    """

    events: list[tuple[float, int, float, float]] = []
    for left, upper, right, lower in rectangles:
        events.append((left, 1, upper, lower))
        events.append((right, -1, upper, lower))
    events.sort()
    active: Counter[tuple[float, float]] = Counter()
    area = 0.0
    previous_x: float | None = None
    index = 0
    while index < len(events):
        current_x = events[index][0]
        if previous_x is not None and current_x > previous_x:
            area += (current_x - previous_x) * _merged_interval_length(active)
        while index < len(events) and events[index][0] == current_x:
            _, change, upper, lower = events[index]
            interval = (upper, lower)
            active[interval] += change
            if active[interval] <= 0:
                del active[interval]
            index += 1
        previous_x = current_x
    return area


def _merged_interval_length(intervals: Counter[tuple[float, float]]) -> float:
    """合并扫描线当前活跃的纵向区间，并返回覆盖总长度。"""

    if not intervals:
        return 0.0
    covered_height = 0.0
    current_upper: float | None = None
    current_lower: float | None = None
    for (upper, lower), count in sorted(intervals.items()):
        if count <= 0:
            continue
        if current_upper is None:
            current_upper, current_lower = upper, lower
        elif current_lower is not None and upper <= current_lower:
            current_lower = max(current_lower, lower)
        else:
            assert current_lower is not None
            covered_height += current_lower - current_upper
            current_upper, current_lower = upper, lower
    if current_upper is not None and current_lower is not None:
        covered_height += current_lower - current_upper
    return covered_height


def _number(value: object) -> float | None:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        return None
    return float(value)


def _pdf_word_coordinate(word: dict[str, object], key: str) -> float:
    """读取已通过预筛选的 PDF word 坐标，并让类型检查器看到数值边界。"""

    value = _number(word.get(key))
    if value is None:
        # 调用方只传入了 _group_pdf_words_into_lines 的有效 word；这里的异常表示
        # 第三方解析库返回了不符合契约的数据，宁可跳过当前解析也不猜坐标。
        raise ValueError(f"PDF word coordinate {key} is not numeric")
    return value


def _positive_number(value: object) -> float:
    number = _number(value)
    return number if number is not None and number > 0 else 0.0


def _area_ratio(area: float, page_area: float) -> float:
    if page_area <= 0:
        return 0.0
    return min(max(area / page_area, 0.0), 1.0)


def _normalize_pdf_line(line: str) -> str:
    """修复 PDF 常见的兼容字符、重复字形和空白，但不改动正文语义。

    某些演示文稿 PDF 的字体映射会把每个 ASCII 字符提取两次，例如把
    ``Information`` 提取为 ``IInnffoorrmmaattiioonn``。不能直接把所有相邻重复
    字符压缩，否则 ``coffee``、``letter`` 等正常单词会被破坏。因此这里只在一个
    完整单词的每一对字符都完全相同时才折叠，例如 ``IInn`` → ``In``、``22nndd``
    → ``2nd``；普通英文单词保持不变。
    """

    normalized = unicodedata.normalize("NFKC", line).replace("\u00a0", " ")
    # 某些 HHS PDF 把项目符号映射为 C1 控制区字符 U+0083。保留为标准项目符号，
    # 让后续父子切分知道这是列表项，同时避免质量报告把它当作不可见脏字符。
    normalized = normalized.replace("\x83", "•")
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return _repair_pdf_duplicate_glyphs(normalized)


_PDF_ASCII_TOKEN = re.compile(r"[A-Za-z0-9]+")


def _repair_pdf_duplicate_glyphs(line: str) -> str:
    """只折叠“整词每个字形都重复”的 PDF 提取噪声。"""

    changed = False

    def replace_token(match: re.Match[str]) -> str:
        nonlocal changed
        token = match.group(0)
        if len(token) < 4 or len(token) % 2 != 0:
            return token
        if all(token[index] == token[index + 1] for index in range(0, len(token), 2)):
            changed = True
            return token[::2]
        return token

    repaired = _PDF_ASCII_TOKEN.sub(replace_token, line)
    if not changed:
        return repaired
    # 字体映射异常可能同时把逗号、句号和 URL 路径分隔符复制一遍。只有同一行已经
    # 证实存在“整词重复字形”时才做这些窄规则，避免把正常省略号、网址协议或英文
    # 双写字母误判成噪声。
    repaired = re.sub(r"(?<!\.)\.\.(?!\.)", ".", repaired)
    repaired = re.sub(r"(?<!,),,(?!,)", ",", repaired)
    repaired = re.sub(r"(?<=[A-Za-z0-9])//(?=[A-Za-z0-9])", "/", repaired)
    return repaired


def _is_pdf_template_line(line: str) -> bool:
    """判断网页打印模板行，规则只删除高置信度的导航/版权内容。"""

    return any(pattern.search(line) for pattern in _PDF_TEMPLATE_PATTERNS)


def _is_pdf_layout_noise_table(table: Sequence[Sequence[str | None]]) -> bool:
    """过滤网页导航和误识别布局表格，但保留正文业务表格。

    PDF 中的多栏文本框有时会被表格算法误判成“一列多行文本表格”。这类对象没有
    表头、没有行列语义，直接走表格提取反而容易重复字符；把它交回正文文本层更可靠。
    """

    raw_cells = [
        str(cell) for row in table for cell in row if cell is not None and str(cell).strip()
    ]
    if len(raw_cells) == 1 and "\n" in raw_cells[0]:
        return True
    cells = [
        _normalize_pdf_line(str(cell))
        for row in table
        for cell in row
        if cell is not None and str(cell).strip()
    ]
    if not cells:
        return True
    joined = " ".join(cells)
    return any(
        marker in joined
        for marker in (
            "无障碍浏览",
            "网站导航",
            "公务员邮箱",
            "请输入关键字",
            "首页 >",
            "打印此页",
            "关闭窗口",
            "版权所有",
        )
    )


def _repeated_pdf_lines(page_lines: Sequence[Sequence[str]]) -> frozenset[str]:
    """识别出现在页首/页尾且跨页重复的页眉页脚，避免误删正文标题。"""

    counts: Counter[str] = Counter()
    edge_counts: Counter[str] = Counter()
    for lines in page_lines:
        normalized_lines = [_normalize_pdf_line(line) for line in lines if line.strip()]
        counts.update(normalized_lines)
        edge_lines = normalized_lines[:3] + normalized_lines[-3:]
        edge_counts.update(edge_lines)
    return frozenset(
        line
        for line, count in counts.items()
        if count >= 2
        and edge_counts[line] >= 2
        and len(line) <= 180
        and line[-1:] not in _PDF_SENTENCE_END
    )


def _clean_pdf_lines(
    raw_lines: Sequence[str],
    *,
    repeated_lines: frozenset[str] = frozenset(),
) -> list[str]:
    """清除高置信度模板行、目录页码点线，并统一 PDF 字体映射字符。"""

    cleaned: list[str] = []
    for raw_line in raw_lines:
        line = _normalize_pdf_line(raw_line)
        if not line or _is_pdf_template_line(line):
            continue
        if line in repeated_lines or _PDF_DOT_LEADER.search(line):
            continue
        cleaned.append(line)
    return cleaned


def _is_pdf_toc_page(lines: Sequence[str]) -> bool:
    """检测目录页并整页排除，避免只删除页码后留下无语义的标题列表。"""

    lowered_lines = [line.lower().strip() for line in lines[:12]]
    has_toc_marker = any(
        line == marker or line == "contents" or line.startswith(f"{marker} ")
        for line in lowered_lines
        for marker in _PDF_TOC_MARKERS
    )
    return has_toc_marker and len(lines) >= 3


def _is_pdf_structural_line(line: str) -> bool:
    """识别标题、列表和编号项，用于限制父节点范围。"""

    if line.startswith(("•", "▪", "- ", "* ")):
        return True
    if re.match(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）|\d+[、.)])", line):
        return True
    if line.endswith(("？", "!", "?", "!")):
        return True
    if re.match(r"^(?:第\s*\S+章|Chapter\s+\d+|Appendix\s+\d+)", line, re.IGNORECASE):
        return True
    if _CJK.search(line):
        # 中文正文在 PDF 中常被按视觉宽度拆成多行，不能仅凭“短”就切开；
        # 只有很短且没有逗号/冒号等连接标点的行，才保守地视为小标题。
        return (
            len(line) <= 16
            and not any(mark in line for mark in "，,；;：:、")
            and not line.endswith(tuple(_PDF_SENTENCE_END))
        )
    # 英文 PDF 也会把一个段落按版心宽度拆成多行，因此不能把所有短行都当标题。
    # 仅接受标题式大小写、明确的章节名或版本名，降低普通正文被切碎的概率。
    if re.match(r"^(?:\d+(?:st|nd|rd|th)?\s+edition)$", line, re.IGNORECASE):
        return True
    words = re.findall(r"[A-Za-z][A-Za-z'-]*", line)
    if not words or len(line) > 64 or len(words) > 8:
        return False
    # 正式指南和 WHO 文档大量使用全大写章节标题，例如“MESSAGE FROM THE
    # SECRETARY”。这类行通常是结构边界，不应按短正文碎片统计；限制在 8 个词以内，
    # 避免把整段大写正文或长表格误当成标题。
    if len(words) >= 2 and all(word.isupper() for word in words):
        return True
    return all(word[0].isupper() or word.lower() in _PDF_ENGLISH_HEADING_WORDS for word in words)


def _is_pdf_section_heading(line: str) -> bool:
    """判断一行是否是段落边界标题，而不是需要继续拼接的列表项。"""

    # 列表项本身要开启新块，但它的下一行可能是同一个项目的续行，不能因为上一行
    # 是结构行就立即 flush；否则英文断词和中文长列表会被切成大量碎片。
    if line.startswith(("•", "▪", "- ", "* ")):
        return False
    if re.match(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）)", line):
        return True
    # 中文资料常用“1、心肺功能”表示章节小标题，而动作步骤更常用“1. 两脚
    # 分开”。只把前者视为标题，避免把健身动作步骤错误升级成章节层级。
    if re.match(r"^\d+、", line):
        return len(line) <= 32 and not any(mark in line for mark in "，,；;：:")
    if re.match(r"^\d+[.)]", line):
        return False
    return _is_pdf_structural_line(line)


def _join_pdf_lines(lines: Sequence[str]) -> str:
    """按中英文边界合并 PDF 换行，避免中文被插入空格、英文单词被粘连。"""

    result = ""
    for line in lines:
        if not result:
            result = line
            continue
        previous = result[-1:]
        first = line[:1]
        if previous == "-" and first.isascii() and first.isalpha():
            # 断词修复不能把真实复合词改成不存在的单词，例如
            # ``fall-`` + ``related`` 必须保留为 ``fall-related``；而
            # ``physical-`` + ``activity`` 通常是排版换行，应合并为
            # ``physicalactivity``。正式项目中可再接领域词典，这里先对白名单复合词
            # 保守保留连字符，其他情况按“行尾连字符是排版符号”处理。
            previous_word = re.search(r"([A-Za-z][A-Za-z-]*)-$", result)
            next_word = re.match(r"([A-Za-z]+)", line)
            compound = (
                f"{previous_word.group(1)}-{next_word.group(1)}".lower()
                if previous_word and next_word
                else ""
            )
            if compound not in _PDF_COMMON_COMPOUNDS:
                result = result[:-1] + line
            else:
                result += line
        elif (
            (_CJK.search(previous) and _CJK.search(first))
            or (previous in "。！？；：，," and _CJK.search(first))
            or previous in "([{／/"
            or first in "，。；：！？、)]}】》%％,.!?;:/"
        ):
            result += line
        else:
            result += " " + line
    return result.strip()


def _split_pdf_text_blocks(lines: Sequence[str], *, max_block_chars: int = 900) -> list[str]:
    """将页面文本切成有界上下文块，避免父节点直接膨胀为整页。"""

    blocks: list[str] = []
    current: list[str] = []
    current_length = 0

    def flush() -> None:
        nonlocal current, current_length
        if current:
            content = _join_pdf_lines(current)
            if content:
                blocks.append(content)
        current = []
        current_length = 0

    for line in lines:
        if current and (_is_pdf_section_heading(current[-1]) or _is_pdf_structural_line(line)):
            flush()
        current.append(line)
        current_length += len(line)
        if (current_length >= max_block_chars and line[-1:] in _PDF_SENTENCE_END) or (
            current_length >= max_block_chars + 180
        ):
            flush()
    flush()
    return blocks


def _pdf_text_blocks_with_context(
    lines: Sequence[str],
) -> list[tuple[str, tuple[str, ...], str]]:
    """按 PDF 标题边界生成正文块，并为同一章节建立共享父上下文。

    PDF 没有 Markdown 的标题语法，不能直接复用 ``chunk_markdown``。这里使用已经
    通过页面清洗的标题/编号启发式：标题进入 ``heading_path``，后续正文按章节聚合，
    最终由子块切分器控制 Embedding 大小。父上下文保留整段章节（而不是只保留当前
    900 字块），这样检索命中后仍能恢复动作名称、适用人群和安全前提。
    """

    result: list[tuple[str, tuple[str, ...], str]] = []
    heading_path: list[str] = []
    section_lines: list[str] = []

    def flush_section() -> None:
        if not section_lines:
            return
        parent_body = _join_pdf_lines(section_lines)
        parent_prefix = " / ".join(heading_path)
        parent_content = f"{parent_prefix}\n{parent_body}" if parent_prefix else parent_body
        for content in _split_pdf_text_blocks(section_lines):
            if content:
                result.append((content, tuple(heading_path), parent_content))
        section_lines.clear()

    for line in lines:
        if _is_pdf_section_heading(line):
            flush_section()
            level = _pdf_heading_level(line)
            heading_path[:] = heading_path[: level - 1]
            heading_path.append(line)
            # 标题本身也要保留。标题块不再把自己放入 heading_path，否则分块器会把
            # 标题前缀和标题正文复制一遍；后续正文块才继承完整的章节路径。
            result.append((line, (), line))
            continue
        section_lines.append(line)
    flush_section()
    return result


def _pdf_heading_level(line: str) -> int:
    """把常见中文/英文章节编号映射到有限的父节点层级。"""

    if re.match(r"^(?:第\s*\S+章|Chapter\s+\d+|Appendix\s+\d+)", line, re.IGNORECASE):
        return 1
    if re.match(r"^[一二三四五六七八九十]+、", line):
        return 1
    if re.match(r"^（[一二三四五六七八九十]+）", line):
        return 2
    if re.match(r"^\d+[、.)]", line):
        return 3
    return 1


def _extract_pdf_text_lines(
    page: object,
    tables: Sequence[Sequence[Sequence[str | None]]],
) -> tuple[list[str], int]:
    """提取页面文字，并在确有多区域版式时按坐标恢复阅读顺序。

    处理策略是“默认不猜，证据充分才重排”：

    * 普通单栏正文沿用 pdfplumber 的原生文本流，降低误重排正文的风险；
    * 只有检测到至少三个同时存在的横向区域，才将文字按区域从左到右、区域内从
      上到下排序，主要解决 PPT 流程图、健身动作图解和复杂双栏页面；
    * 表格区域中的文字已经单独转成 Markdown，不再混入正文，避免同一内容重复两次；
    * 坐标提取失败时安全回退到原有实现，不能因为版面增强导致整份文档不可用。
    """

    default_text = _extract_pdf_text_outside_tables(page, tables)
    default_lines = [line for line in default_text.splitlines() if line.strip()]
    try:
        words = page.extract_words()  # type: ignore[attr-defined]
    except Exception:  # noqa: BLE001 - 坐标不是所有 PDF 都可靠，回退到原生文本流
        return default_lines, 1
    if not isinstance(words, Sequence):
        return default_lines, 1

    table_boxes = _pdf_table_boxes(page, tables)
    filtered_words = [
        word
        for word in words
        if isinstance(word, dict) and not _pdf_word_inside_boxes(word, table_boxes)
    ]
    coordinate_lines = _group_pdf_words_into_lines(filtered_words)
    column_count = _detect_pdf_column_count(
        coordinate_lines, _positive_number(getattr(page, "width", 0.0))
    )
    # 坐标重排是高风险操作：普通双栏正文通常已经被 pdfplumber 正确还原，而图解
    # 页面才经常把多个独立文本框交叉拼接。因此默认只对图片占比很高、且同时有
    # 三个以上横向区域的页面启用，避免把正常指南正文“优化”坏。
    image_area_ratio = _area_ratio(
        _rectangle_union_area(
            _pdf_object_rectangles(
                getattr(page, "images", ()),
                _positive_number(getattr(page, "width", 0.0)),
                _positive_number(getattr(page, "height", 0.0)),
            )
        ),
        _positive_number(getattr(page, "width", 0.0))
        * _positive_number(getattr(page, "height", 0.0)),
    )
    if (
        image_area_ratio < 0.45
        or column_count < 3
        or len(coordinate_lines) < max(3, len(default_lines) // 2)
    ):
        return default_lines, 1

    ordered = _order_pdf_layout_lines(
        coordinate_lines,
        page_width=_positive_number(getattr(page, "width", 0.0)),
    )
    return [line.text for line in ordered if line.text.strip()], column_count


def _pdf_table_boxes(
    page: object,
    tables: Sequence[Sequence[Sequence[str | None]]],
) -> list[tuple[float, float, float, float]]:
    """取得带坐标表格的区域；回退表格没有坐标时返回空列表。"""

    if not tables:
        return []
    try:
        table_objects = page.find_tables()  # type: ignore[attr-defined]
    except Exception:  # noqa: BLE001
        return []
    boxes: list[tuple[float, float, float, float]] = []
    for table in table_objects:
        bbox = getattr(table, "bbox", None)
        if (
            isinstance(bbox, (tuple, list))
            and len(bbox) == 4
            and all(isinstance(value, (int, float)) for value in bbox)
        ):
            boxes.append(tuple(float(value) for value in bbox))  # type: ignore[arg-type]
    return boxes


def _pdf_word_inside_boxes(
    word: dict[str, object], boxes: Sequence[tuple[float, float, float, float]]
) -> bool:
    """按文字中心点排除表格区域，避免误删跨越表格边界的正文。"""

    x0, x1 = _number(word.get("x0")), _number(word.get("x1"))
    top, bottom = _number(word.get("top")), _number(word.get("bottom"))
    if None in (x0, x1, top, bottom):
        return False
    assert x0 is not None and x1 is not None and top is not None and bottom is not None
    center_x, center_y = (x0 + x1) / 2, (top + bottom) / 2
    return any(
        left <= center_x <= right and upper <= center_y <= lower
        for left, upper, right, lower in boxes
    )


def _group_pdf_words_into_lines(words: Sequence[object]) -> list[_PdfTextLine]:
    """按纵向基线和横向大间隙把 pdfplumber words 聚合成可排序文字行。"""

    valid_words = [
        word
        for word in words
        if isinstance(word, dict)
        and isinstance(word.get("text"), str)
        and _number(word.get("x0")) is not None
        and _number(word.get("x1")) is not None
        and _number(word.get("top")) is not None
        and _number(word.get("bottom")) is not None
    ]
    if not valid_words:
        return []
    heights = [
        max(
            0.1,
            _pdf_word_coordinate(word, "bottom") - _pdf_word_coordinate(word, "top"),
        )
        for word in valid_words
        if isinstance(word["top"], (int, float)) and isinstance(word["bottom"], (int, float))
    ]
    line_tolerance = max(2.0, (sum(heights) / len(heights)) * 0.35)
    page_width = max(_pdf_word_coordinate(word, "x1") for word in valid_words)
    rows: list[list[dict[str, object]]] = []
    for word in sorted(
        valid_words,
        key=lambda item: (
            _pdf_word_coordinate(item, "top"),
            _pdf_word_coordinate(item, "x0"),
        ),
    ):
        target: list[dict[str, object]] | None = None
        for candidate in reversed(rows):
            candidate_top = _pdf_word_coordinate(candidate[0], "top")
            if _pdf_word_coordinate(word, "top") - candidate_top > line_tolerance:
                break
            if abs(_pdf_word_coordinate(word, "top") - candidate_top) <= line_tolerance:
                target = candidate
                break
        if target is None:
            rows.append([word])
        else:
            target.append(word)

    result: list[_PdfTextLine] = []
    for row in rows:
        row.sort(key=lambda item: _pdf_word_coordinate(item, "x0"))
        segments: list[list[dict[str, object]]] = [[]]
        for word in row:
            if segments[-1]:
                gap = _pdf_word_coordinate(word, "x0") - _pdf_word_coordinate(
                    segments[-1][-1], "x1"
                )
                # 同一文本框中的词间距通常远小于字号；明显的大间隙通常意味着
                # 两个独立栏目/图解节点，必须拆开，否则排序前就已被拼坏。
                if gap > max(20.0, page_width * 0.08):
                    segments.append([])
            segments[-1].append(word)
        for segment in segments:
            text = _join_pdf_word_text(segment)
            if text:
                result.append(
                    _PdfTextLine(
                        text=text,
                        x0=_pdf_word_coordinate(segment[0], "x0"),
                        x1=_pdf_word_coordinate(segment[-1], "x1"),
                        top=min(_pdf_word_coordinate(item, "top") for item in segment),
                        bottom=max(_pdf_word_coordinate(item, "bottom") for item in segment),
                    )
                )
    return result


def _join_pdf_word_text(words: Sequence[dict[str, object]]) -> str:
    """按相邻文字的语言边界拼接 words，避免中文词之间出现无意义空格。"""

    result = ""
    previous_x1: float | None = None
    for word in words:
        text = str(word["text"])
        if not result:
            result = text
        else:
            gap = _pdf_word_coordinate(word, "x0") - (
                previous_x1 or _pdf_word_coordinate(word, "x0")
            )
            previous = result[-1:]
            first = text[:1]
            if gap > 1.5 and not (_CJK.search(previous) and _CJK.search(first)):
                result += " "
            result += text
        previous_x1 = _pdf_word_coordinate(word, "x1")
    return result.strip()


def _detect_pdf_column_count(lines: Sequence[_PdfTextLine], page_width: float) -> int:
    """估计页面同时存在的横向区域数量，只用于触发保守重排。"""

    if page_width <= 0 or len(lines) < 6:
        return 1
    clusters = _cluster_pdf_line_starts(lines, page_width=page_width)
    return len(clusters) if len(clusters) >= 3 else 1


def _order_pdf_layout_lines(
    lines: Sequence[_PdfTextLine], *, page_width: float
) -> list[_PdfTextLine]:
    """恢复复杂页面的阅读顺序，并把跨栏标题/页脚放回页面首尾。"""

    if not lines:
        return []
    full_width = [line for line in lines if line.width >= page_width * _PDF_FULL_WIDTH_RATIO]
    body = [line for line in lines if line.width < page_width * _PDF_FULL_WIDTH_RATIO]
    if not body:
        return sorted(lines, key=lambda line: (line.top, line.x0))
    # 跨栏标题一般位于正文前，来源署名/页脚位于正文后；两者不能被塞进某一列中间。
    first_body_top = min(line.top for line in body)
    last_body_bottom = max(line.bottom for line in body)
    prefix = sorted(
        [line for line in full_width if line.bottom <= first_body_top + 3],
        key=lambda line: line.top,
    )
    suffix = sorted(
        [line for line in full_width if line.top >= last_body_bottom - 3], key=lambda line: line.top
    )
    middle = [line for line in full_width if line not in prefix and line not in suffix]
    ordered_body: list[_PdfTextLine] = []
    columns = _cluster_pdf_line_starts(body, page_width=page_width)
    centers = [sum(column) / len(column) for column in columns]
    for column_start in centers:
        column_lines = [
            line
            for line in body
            if min(abs(line.x0 - center) for center in centers) == abs(line.x0 - column_start)
        ]
        ordered_body.extend(sorted(column_lines, key=lambda line: (line.top, line.x0)))
    return [*prefix, *ordered_body, *sorted(middle, key=lambda line: (line.top, line.x0)), *suffix]


def _cluster_pdf_line_starts(
    lines: Sequence[_PdfTextLine], *, page_width: float
) -> list[list[float]]:
    """把有缩进差异的行首聚成版面区域，而不是把每个缩进当作新栏目。"""

    starts = sorted(
        {round(line.x0, 1) for line in lines if line.width < page_width * _PDF_FULL_WIDTH_RATIO}
    )
    if not starts:
        return []
    gap_threshold = max(30.0, page_width * 0.08)
    clusters: list[list[float]] = [[starts[0]]]
    for start in starts[1:]:
        if start - clusters[-1][-1] > gap_threshold:
            clusters.append([])
        clusters[-1].append(start)
    return clusters


def _count_pdf_dehyphenations(lines: Sequence[str]) -> int:
    """统计将要在行合并阶段修复的英文断词数量。"""

    return sum(
        1
        for previous, current in pairwise(lines)
        if previous.endswith("-") and current[:1].isascii() and current[:1].isalpha()
    )


def _extract_pdf_tables(page: object) -> list[list[list[str | None]]]:
    """优先使用带坐标的表格对象，为正文排除表格区域并保留表头。"""

    try:
        tables = page.find_tables()  # type: ignore[attr-defined]
        extracted = [table.extract() for table in tables]
    except Exception:  # noqa: BLE001 - 单页表格失败时回退到 pdfplumber 基础接口
        extracted = []
    extracted = [table for table in extracted if not _is_pdf_layout_noise_table(table)]
    if extracted:
        return extracted
    try:
        fallback = page.extract_tables() or []  # type: ignore[attr-defined]
        return [table for table in fallback if not _is_pdf_layout_noise_table(table)]
    except Exception:  # noqa: BLE001 - 解析失败由空表交给正文/OCR流程处理
        return []


def _extract_pdf_text_outside_tables(
    page: object,
    tables: Sequence[Sequence[Sequence[str | None]]],
) -> str:
    """提取页面正文；表格坐标可用时先从文本层排除表格，避免表格重复入库。"""

    # ``extract_tables`` 的回退结果没有坐标，不能猜测区域；此时保留文本并依赖后续
    # 质量评测发现重复，避免错误裁剪正文。
    try:
        table_objects = page.find_tables()  # type: ignore[attr-defined]
    except Exception:  # noqa: BLE001
        table_objects = []
    if not table_objects or not tables:
        return page.extract_text(layout=False) or ""  # type: ignore[attr-defined]
    boxes = [table.bbox for table in table_objects]

    def keep_object(obj: dict[str, object]) -> bool:
        x0, x1 = obj.get("x0"), obj.get("x1")
        top, bottom = obj.get("top"), obj.get("bottom")
        if not all(isinstance(value, (int, float)) for value in (x0, x1, top, bottom)):
            return True
        assert isinstance(x0, (int, float))
        assert isinstance(x1, (int, float))
        assert isinstance(top, (int, float))
        assert isinstance(bottom, (int, float))
        center_x = (x0 + x1) / 2
        center_y = (top + bottom) / 2
        return not any(
            left <= center_x <= right and upper <= center_y <= lower
            for left, upper, right, lower in boxes
        )

    try:
        filtered_page = page.filter(keep_object)  # type: ignore[attr-defined]
        return filtered_page.extract_text(layout=False) or ""
    except Exception:  # noqa: BLE001 - 坐标过滤失败时保留原文本，不丢正文
        return page.extract_text(layout=False) or ""  # type: ignore[attr-defined]


def _table_rows_for_metadata(
    rows: Sequence[Sequence[str | None]],
) -> list[list[str]]:
    """按与 Markdown 渲染相同的规则规范化表格行，供续表判断使用。"""

    cleaned = [
        ["" if cell is None else _normalize_pdf_line(str(cell)) for cell in row] for row in rows
    ]
    return [row for row in cleaned if any(cell.strip() for cell in row)]


def _table_column_count(rows: Sequence[Sequence[str | None]]) -> int:
    """返回规范化表格的最大列数。"""

    return max((len(row) for row in _table_rows_for_metadata(rows)), default=0)


def _table_row_count(rows: Sequence[Sequence[str | None]]) -> int:
    """返回规范化表格的非空行数，包含表头。"""

    return len(_table_rows_for_metadata(rows))


def _table_header_signature(rows: Sequence[Sequence[str | None]]) -> str:
    """返回去空白、大小写无关的完整表头签名。"""

    normalized = _table_rows_for_metadata(rows)
    if not normalized:
        return ""
    return "|".join(cell.casefold() for cell in normalized[0])


def _table_header_key(rows: Sequence[Sequence[str | None]]) -> str:
    """返回第一列表头，用于发现疑似续表但列形状不一致的情况。"""

    normalized = _table_rows_for_metadata(rows)
    return normalized[0][0].casefold() if normalized and normalized[0] else ""


def _annotate_pdf_table_continuations(blocks: Sequence[ParsedBlock]) -> list[ParsedBlock]:
    """为相邻页面的表格建立保守续接关系，不直接拼接不确定的表格。"""

    table_positions = [
        index
        for index, block in enumerate(blocks)
        if block.kind == "TABLE" and block.source_page is not None
    ]
    if not table_positions:
        return list(blocks)

    tables_by_page: dict[int, list[int]] = {}
    for position in table_positions:
        page = blocks[position].source_page
        assert page is not None
        tables_by_page.setdefault(page, []).append(position)

    annotated = list(blocks)
    # 先为每张表设置稳定的单页组标识；只有满足强证据时才共享组标识。
    for position in table_positions:
        block = annotated[position]
        metadata = dict(block.metadata or {})
        page = block.source_page
        table_index = block.table_index
        metadata.setdefault(
            "table_continuation_group",
            f"page-{page}-table-{table_index if table_index is not None else position}",
        )
        metadata.setdefault("table_continuation_status", "SINGLE_PAGE")
        annotated[position] = replace(block, metadata=metadata)

    for page in sorted(tables_by_page):
        next_page = page + 1
        if next_page not in tables_by_page:
            continue
        previous_positions = tables_by_page[page]
        next_positions = tables_by_page[next_page]
        for previous_position in previous_positions:
            previous = annotated[previous_position]
            previous_metadata = previous.metadata or {}
            for next_position in next_positions:
                current = annotated[next_position]
                current_metadata = current.metadata or {}
                same_header = bool(
                    previous_metadata.get("table_header_signature")
                    and previous_metadata.get("table_header_signature")
                    == current_metadata.get("table_header_signature")
                )
                same_columns = previous_metadata.get("table_column_count") == current_metadata.get(
                    "table_column_count"
                )
                same_first_header = bool(
                    previous_metadata.get("table_header_key")
                    and previous_metadata.get("table_header_key")
                    == current_metadata.get("table_header_key")
                )
                if same_header and same_columns:
                    if len(previous_positions) == len(next_positions) == 1:
                        group_id = str(previous_metadata["table_continuation_group"])
                        previous_metadata = {
                            **previous_metadata,
                            "table_continuation_group": group_id,
                            "table_continuation_status": "CONTINUATION_START",
                        }
                        current_metadata = {
                            **current_metadata,
                            "table_continuation_group": group_id,
                            "table_continuation_status": "CONTINUATION",
                        }
                        annotated[previous_position] = replace(
                            previous,
                            metadata=previous_metadata,
                        )
                        annotated[next_position] = replace(current, metadata=current_metadata)
                    else:
                        # 多表页面即使表头相同，也无法仅凭页码确认对应关系，必须人工复核。
                        for candidate_position in (previous_position, next_position):
                            candidate = annotated[candidate_position]
                            candidate_metadata = {
                                **(candidate.metadata or {}),
                                "table_continuation_status": "AMBIGUOUS_REVIEW",
                            }
                            annotated[candidate_position] = replace(
                                candidate,
                                metadata=candidate_metadata,
                            )
                elif same_first_header and not same_columns:
                    # 第一列表头相同但列数变化，可能是续表被截断，也可能是相邻的另一张表。
                    # 不拼接，统一显式标记形状不一致，交给质量门禁/人工复核。
                    for candidate_position in (previous_position, next_position):
                        candidate = annotated[candidate_position]
                        candidate_metadata = {
                            **(candidate.metadata or {}),
                            "table_continuation_status": "SHAPE_MISMATCH_REVIEW",
                        }
                        annotated[candidate_position] = replace(
                            candidate,
                            metadata=candidate_metadata,
                        )
    return annotated


class DocxParser:
    """按 DOCX 文档顺序读取段落和表格，并保留标题上下文。"""

    def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        from docx import Document
        from docx.document import Document as DocumentType
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document: DocumentType = Document(BytesIO(content))
        blocks: list[ParsedBlock] = []
        heading_path: list[str] = []
        table_index = 0
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = Paragraph(child, document)
                text = " ".join(paragraph.text.split()).strip()
                if not text:
                    continue
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name.lower().startswith("heading"):
                    level = _heading_level(style_name)
                    heading_path[:] = heading_path[: level - 1]
                    heading_path.append(text)
                    continue
                blocks.append(
                    ParsedBlock(kind="TEXT", content=text, heading_path=tuple(heading_path))
                )
            elif child.tag.endswith("}tbl"):
                table = Table(child, document)
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                normalized = _table_to_markdown(rows)
                if normalized:
                    blocks.append(
                        ParsedBlock(
                            kind="TABLE",
                            content=normalized,
                            heading_path=tuple(heading_path),
                            table_index=table_index,
                        )
                    )
                    table_index += 1
        if not blocks:
            raise DocumentParseError("DOCX contains no indexable paragraphs or tables")
        return ParsedDocument(
            tuple(blocks), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


class XlsxParser:
    """将非空工作表转换为保留表头的表格内容块。"""

    def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        from openpyxl import load_workbook  # type: ignore[import-untyped]

        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        blocks: list[ParsedBlock] = []
        warnings: list[str] = []
        try:
            for sheet in workbook.worksheets:
                rows = [
                    ["" if value is None else str(value).strip() for value in row]
                    for row in sheet.iter_rows(values_only=True)
                ]
                rows = [row for row in rows if any(cell for cell in row)]
                if not rows:
                    warnings.append(f"worksheet {sheet.title} is empty")
                    continue
                normalized = _table_to_markdown(rows)
                if normalized:
                    blocks.append(
                        ParsedBlock(
                            kind="TABLE",
                            content=normalized,
                            heading_path=(sheet.title,),
                            source_sheet=sheet.title,
                            table_index=len(blocks),
                            metadata={"parser": "openpyxl", "worksheet": sheet.title},
                        )
                    )
        finally:
            workbook.close()
        if not blocks:
            raise DocumentParseError("XLSX contains no non-empty worksheets")
        return ParsedDocument(
            tuple(blocks),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            tuple(warnings),
        )


def _heading_level(style_name: str) -> int:
    """读取 DOCX Heading 1..6 样式，对未知标题样式进行安全限制。"""

    match = next((char for char in style_name if char.isdigit()), "1")
    return max(1, min(6, int(match)))


def _table_to_markdown(rows: Sequence[Sequence[str | None]]) -> str:
    """渲染表格并重复表头，使按行切分的内容块仍能自描述。"""

    cleaned_rows = [
        ["" if cell is None else _escape_cell(_normalize_pdf_line(str(cell))) for cell in row]
        for row in rows
    ]
    cleaned_rows = [row for row in cleaned_rows if any(cell.strip() for cell in row)]
    if not cleaned_rows:
        return ""
    width = max(len(row) for row in cleaned_rows)
    padded = [row + [""] * (width - len(row)) for row in cleaned_rows]
    header = padded[0]
    separator = ["---"] * width
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(separator) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in padded[1:])
    return "\n".join(lines)


def _escape_cell(value: str) -> str:
    """单元格包含 ``|`` 时，仍保持管道分隔表格结构完整。"""

    return " ".join(value.replace("|", "\\|").split())
