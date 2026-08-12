from io import BytesIO

import pytest

from app.rag.formats import (
    DocumentParseError,
    DocumentParserRegistry,
    PdfPageRoutingPolicy,
    UnsupportedDocumentFormatError,
    _clean_pdf_lines,
    _is_pdf_layout_noise_table,
    _join_pdf_lines,
    _profile_pdf_page,
    _rectangle_union_area,
    _repeated_pdf_lines,
    _split_pdf_text_blocks,
)


def test_registry_parses_markdown_and_rejects_unknown_extensions() -> None:
    registry = DocumentParserRegistry(max_source_bytes=1000)

    parsed = registry.parse(b"# Warmup\n\nPrepare the hips.", file_name="guide.md")

    assert parsed.media_type == "text/markdown"
    assert parsed.blocks[0].content == "# Warmup\n\nPrepare the hips."
    with pytest.raises(UnsupportedDocumentFormatError):
        registry.parse(b"content", file_name="guide.rtf")


def test_registry_enforces_source_size_before_parser_work() -> None:
    registry = DocumentParserRegistry(max_source_bytes=3)

    with pytest.raises(DocumentParseError, match="size limit"):
        registry.parse(b"abcd", file_name="guide.txt")


def test_docx_parser_preserves_heading_and_table_structure() -> None:
    from docx import Document

    document = Document()
    document.add_heading("Strength Training", level=1)
    document.add_paragraph("Use a controlled range of motion.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Exercise"
    table.rows[0].cells[1].text = "Sets"
    row = table.add_row()
    row.cells[0].text = "Squat"
    row.cells[1].text = "4"
    payload = BytesIO()
    document.save(payload)

    parsed = DocumentParserRegistry().parse(payload.getvalue(), file_name="guide.docx")

    assert parsed.blocks[0].heading_path == ("Strength Training",)
    assert parsed.blocks[0].content == "Use a controlled range of motion."
    assert parsed.blocks[1].kind == "TABLE"
    assert parsed.blocks[1].heading_path == ("Strength Training",)
    assert "| Exercise | Sets |" in parsed.blocks[1].content
    assert "| Squat | 4 |" in parsed.blocks[1].content


def test_xlsx_parser_keeps_worksheet_as_a_header_preserving_table() -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Beginner Plan"
    sheet.append(["Exercise", "Sets", "Reps"])
    sheet.append(["Squat", 4, 12])
    payload = BytesIO()
    workbook.save(payload)

    parsed = DocumentParserRegistry().parse(payload.getvalue(), file_name="plan.xlsx")

    block = parsed.blocks[0]
    assert block.kind == "TABLE"
    assert block.source_sheet == "Beginner Plan"
    assert block.heading_path == ("Beginner Plan",)
    assert "| Exercise | Sets | Reps |" in block.content
    assert "| Squat | 4 | 12 |" in block.content


def test_pdf_parser_preserves_ocr_route_for_scanned_pdf_without_linux_service() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    payload = BytesIO()
    writer.write(payload)

    parsed = DocumentParserRegistry().parse(payload.getvalue(), file_name="scan.pdf")

    assert parsed.blocks == ()
    assert parsed.page_profiles[0].route == "OCR_REQUIRED"
    assert "requires OCR" in parsed.warnings[0]


def test_pdf_cleaning_normalizes_font_mapping_and_removes_web_templates() -> None:
    lines = _clean_pdf_lines(
        [
            "⽆障碍浏览 ⽹站导航 公务员邮箱",
            "⾸ 页 机 构 公 开 资 讯 服 务 互 动",
            "⾼温运动防护四注意",
            "【打印此页】 【关闭窗⼜】",
            "⼀、科学补⽔：运动前30分钟喝200⾄300毫升温⽔。",
        ]
    )

    assert lines == ["高温运动防护四注意", "一、科学补水:运动前30分钟喝200至300毫升温水。"]


def test_pdf_cleaning_removes_toc_page_and_repeated_short_headers() -> None:
    repeated = frozenset({"Physical Activity Guidelines for Americans"})
    lines = _clean_pdf_lines(
        [
            "Physical Activity Guidelines for Americans",
            "Table of Contents",
            "Chapter 1 ........................................ 13",
        ],
        repeated_lines=repeated,
    )

    assert lines == ["Table of Contents"]


def test_pdf_cleaning_removes_repeated_chinese_edge_headers_only() -> None:
    repeated = _repeated_pdf_lines(
        [
            ["全民健身指南", "第一页正文内容。"],
            ["全民健身指南", "第二页正文内容。"],
        ]
    )

    assert _clean_pdf_lines(["全民健身指南", "正文标题"], repeated_lines=repeated) == ["正文标题"]


def test_pdf_cleaning_filters_web_layout_tables_but_keeps_domain_tables() -> None:
    assert _is_pdf_layout_noise_table([["无障碍浏览", "网站导航"], ["公务员邮箱", "请输入关键字"]])
    assert not _is_pdf_layout_noise_table([["动作", "组数", "次数"], ["深蹲", "4", "12"]])


def test_pdf_text_blocks_bound_parent_context_and_preserve_chinese_spacing() -> None:
    blocks = _split_pdf_text_blocks(
        [
            "如何提高肌肉耐力？",
            "有氧运动可提升心肺功能，优化肌肉摄氧效率。",
            "可以通过慢跑、游泳、骑自行车等运动来实现。",
        ],
        max_block_chars=100,
    )

    assert blocks == [
        "如何提高肌肉耐力？",
        "有氧运动可提升心肺功能，优化肌肉摄氧效率。可以通过慢跑、游泳、骑自行车等运动来实现。",
    ]
    assert _join_pdf_lines(["Physical Activity", "Guidelines"]) == "Physical Activity Guidelines"


def test_pdf_text_blocks_do_not_split_chinese_sentence_fragments_by_length() -> None:
    blocks = _split_pdf_text_blocks(
        [
            "第一段正文在 PDF 中被换行，",
            "后续内容仍然属于同一个段落并且没有结束标点",
            "第二段标题",
            "第二段正文。",
        ]
    )

    assert blocks == [
        "第一段正文在 PDF 中被换行，后续内容仍然属于同一个段落并且没有结束标点",
        "第二段标题",
        "第二段正文。",
    ]


class FakePdfPage:
    """只暴露页面画像所需坐标，避免单元测试依赖 OCR 或操作系统。"""

    width = 100
    height = 100

    def __init__(self, *, images: list[dict[str, int]], chars: list[dict[str, int]]) -> None:
        self.images = images
        self.chars = chars


def _rect(x0: int, top: int, x1: int, bottom: int) -> dict[str, int]:
    return {"x0": x0, "top": top, "x1": x1, "bottom": bottom}


def test_pdf_page_profile_routes_image_heavy_fitness_page_to_professional_review() -> None:
    page = FakePdfPage(
        images=[_rect(5, 5, 95, 80)],
        chars=[_rect(10, 82, 12, 85), _rect(13, 82, 15, 85)],
    )

    profile = _profile_pdf_page(
        page,
        page_number=3,
        raw_text="图3 深蹲动作示意，保持膝盖与脚尖方向一致。",
        tables=(),
        policy=PdfPageRoutingPolicy(),
    )

    assert profile.route == "VISUAL_REVIEW_REQUIRED"
    assert profile.image_area_ratio == pytest.approx(0.675)
    assert profile.caption_count == 1
    assert "IMAGE_AREA_HIGH" in profile.reasons


def test_pdf_page_profile_routes_low_text_page_to_ocr_without_linux_service() -> None:
    page = FakePdfPage(images=[_rect(0, 0, 100, 100)], chars=[])

    profile = _profile_pdf_page(
        page,
        page_number=1,
        raw_text="",
        tables=(),
        policy=PdfPageRoutingPolicy(),
    )

    assert profile.route == "OCR_REQUIRED"
    assert profile.reasons == ("NATIVE_TEXT_INSUFFICIENT", "PAGE_CONTAINS_IMAGES")


def test_pdf_page_profile_keeps_combined_route_for_short_action_caption() -> None:
    page = FakePdfPage(
        images=[_rect(0, 0, 100, 90)],
        chars=[_rect(10, 92, 15, 95)],
    )

    profile = _profile_pdf_page(
        page,
        page_number=4,
        raw_text="图4 深蹲",
        tables=(),
        policy=PdfPageRoutingPolicy(),
    )

    assert profile.route == "OCR_AND_VISUAL_REVIEW_REQUIRED"
    assert "NATIVE_TEXT_INSUFFICIENT" in profile.reasons
    assert "IMAGE_AREA_HIGH" in profile.reasons


def test_pdf_page_profile_keeps_text_sufficient_page_on_native_path() -> None:
    page = FakePdfPage(images=[_rect(5, 5, 25, 25)], chars=[_rect(10, 30, 90, 80)])

    profile = _profile_pdf_page(
        page,
        page_number=2,
        raw_text="训练前应完成充分热身。" * 80,
        tables=(),
        policy=PdfPageRoutingPolicy(),
    )

    assert profile.route == "NORMAL"


def test_pdf_rectangle_union_does_not_double_count_overlapping_images() -> None:
    area = _rectangle_union_area([(0, 0, 80, 80), (20, 20, 100, 100)])

    assert area == pytest.approx(9_200)
